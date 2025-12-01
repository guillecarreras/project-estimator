"""
DOC/DOCX Parser

Extracts content from Microsoft Word documents while preserving structure and formatting.
"""

from docx import Document
from typing import Dict, List


class DocParser:
    """Parse Word documents and extract content"""

    def __init__(self, console):
        """
        Initialize DOC/DOCX parser

        Args:
            console: Rich console for output
        """
        self.console = console

    def parse(self, file_path: str) -> Dict:
        """
        Parse a Word document and extract its content

        Args:
            file_path: Path to Word document (.doc or .docx)

        Returns:
            dict: Parsed content with structure
        """
        self.console.print(f"[cyan]Parsing Word document: {file_path}[/cyan]")

        try:
            content = {
                'file_path': file_path,
                'file_type': 'DOCX',
                'paragraphs': [],
                'text': '',
                'sections': [],
                'metadata': {},
                'tables': []
            }

            doc = Document(file_path)

            # Extract metadata
            core_props = doc.core_properties
            content['metadata'] = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else ''
            }

            # Extract paragraphs with their styles
            for para in doc.paragraphs:
                para_data = {
                    'text': para.text,
                    'style': para.style.name if para.style else 'Normal',
                    'runs': []
                }

                # Preserve formatting of runs
                for run in para.runs:
                    run_data = {
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline
                    }
                    para_data['runs'].append(run_data)

                content['paragraphs'].append(para_data)

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    'index': table_idx,
                    'rows': []
                }

                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data['rows'].append(row_data)

                content['tables'].append(table_data)

            # Build plain text
            content['text'] = '\n'.join([p['text'] for p in content['paragraphs']])

            # Extract sections based on heading styles
            content['sections'] = self.extract_sections(content['paragraphs'])

            self.console.print(f"[green]✓ Extracted {len(content['paragraphs'])} paragraphs[/green]")
            self.console.print(f"[green]✓ Extracted {len(content['tables'])} tables[/green]")
            self.console.print(f"[green]✓ Detected {len(content['sections'])} sections[/green]")

            return content

        except Exception as e:
            self.console.print(f"[red]Error parsing Word document: {str(e)}[/red]")
            raise

    def extract_sections(self, paragraphs: List[Dict]) -> List[Dict]:
        """
        Extract sections based on heading styles

        Args:
            paragraphs: List of paragraph data

        Returns:
            List of sections with their content
        """
        sections = []
        current_section = None

        for para in paragraphs:
            style = para['style']
            text = para['text']

            # Check if it's a heading
            if 'Heading' in style or 'Title' in style:
                if current_section:
                    sections.append(current_section)

                # Extract heading level
                level = 1
                if 'Heading' in style:
                    try:
                        level = int(''.join(filter(str.isdigit, style)) or '1')
                    except:
                        level = 1

                current_section = {
                    'level': level,
                    'title': text,
                    'style': style,
                    'content': [],
                    'paragraphs': []
                }
            elif current_section:
                current_section['content'].append(text)
                current_section['paragraphs'].append(para)
            else:
                # Content before first heading
                if not sections:
                    sections.append({
                        'level': 0,
                        'title': 'Introduction',
                        'style': 'Normal',
                        'content': [text],
                        'paragraphs': [para]
                    })

        if current_section:
            sections.append(current_section)

        # Join content
        for section in sections:
            section['content'] = '\n'.join(section['content']).strip()

        return sections

    def preserve_formatting(self, paragraphs: List[Dict]) -> str:
        """
        Generate formatted text preserving basic formatting

        Args:
            paragraphs: List of paragraph data

        Returns:
            str: Formatted text with basic markdown-style formatting
        """
        formatted_lines = []

        for para in paragraphs:
            if not para['text'].strip():
                formatted_lines.append('')
                continue

            # Handle headings
            if 'Heading' in para['style']:
                level = int(''.join(filter(str.isdigit, para['style'])) or '1')
                formatted_lines.append('#' * level + ' ' + para['text'])
            else:
                # Reconstruct text with formatting
                line_parts = []
                for run in para['runs']:
                    text = run['text']
                    if run['bold']:
                        text = f"**{text}**"
                    if run['italic']:
                        text = f"*{text}*"
                    line_parts.append(text)

                formatted_lines.append(''.join(line_parts))

        return '\n'.join(formatted_lines)
